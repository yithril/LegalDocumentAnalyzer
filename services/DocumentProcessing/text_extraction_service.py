import os
from abc import ABC, abstractmethod
from typing import Dict, Any
import logging

logger = logging.getLogger(__name__)

class TextExtractionError(Exception):
    """Raised when text extraction fails"""
    pass

class UnsupportedFileTypeError(TextExtractionError):
    """Raised when file type is not supported"""
    pass

class EmptyTextError(TextExtractionError):
    """Raised when text extraction produces empty content"""
    pass

class TextExtractor(ABC):
    """Abstract base class for text extractors"""
    
    @abstractmethod
    def extract(self, file_path: str) -> str:
        """Extract text from file"""
        pass
    
    @abstractmethod
    def get_metadata(self, file_path: str) -> Dict[str, Any]:
        """Get metadata about the extraction process"""
        pass

class TxtExtractor(TextExtractor):
    """Extract text from plain text files"""
    
    def extract(self, file_path: str) -> str:
        try:
            # Try UTF-8 first
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
        except UnicodeDecodeError:
            # Fallback to other encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        text = f.read()
                    break
                except UnicodeDecodeError:
                    continue
            else:
                raise TextExtractionError(f"Could not decode text file: {file_path}")
        
        if not text.strip():
            raise EmptyTextError(f"Text file is empty: {file_path}")
        
        return text
    
    def get_metadata(self, file_path: str) -> Dict[str, Any]:
        return {
            "extraction_method": "plain_text",
            "encoding": "utf-8",  # Simplified for MVP
            "file_size": os.path.getsize(file_path)
        }

class PdfExtractor(TextExtractor):
    """Extract text from PDF files"""
    
    def extract(self, file_path: str) -> str:
        try:
            import PyPDF2
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            if not text.strip():
                raise EmptyTextError(f"PDF contains no extractable text: {file_path}")
            
            return text.strip()
            
        except ImportError:
            raise TextExtractionError("PyPDF2 is required for PDF extraction")
        except Exception as e:
            raise TextExtractionError(f"PDF extraction failed: {str(e)}")
    
    def get_metadata(self, file_path: str) -> Dict[str, Any]:
        try:
            import PyPDF2
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                return {
                    "extraction_method": "pdf_text",
                    "page_count": len(pdf_reader.pages),
                    "file_size": os.path.getsize(file_path)
                }
        except:
            return {
                "extraction_method": "pdf_text",
                "file_size": os.path.getsize(file_path)
            }

class DocxExtractor(TextExtractor):
    """Extract text from DOCX files"""
    
    def extract(self, file_path: str) -> str:
        try:
            from docx import Document
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            if not text.strip():
                raise EmptyTextError(f"DOCX contains no extractable text: {file_path}")
            
            return text.strip()
            
        except ImportError:
            raise TextExtractionError("python-docx is required for DOCX extraction")
        except Exception as e:
            raise TextExtractionError(f"DOCX extraction failed: {str(e)}")
    
    def get_metadata(self, file_path: str) -> Dict[str, Any]:
        try:
            from docx import Document
            doc = Document(file_path)
            return {
                "extraction_method": "docx_text",
                "paragraph_count": len(doc.paragraphs),
                "file_size": os.path.getsize(file_path)
            }
        except:
            return {
                "extraction_method": "docx_text",
                "file_size": os.path.getsize(file_path)
            }

class TextExtractionService:
    """Main service for text extraction using strategy pattern"""
    
    def __init__(self):
        self.extractors = {
            '.txt': TxtExtractor(),
            '.pdf': PdfExtractor(),
            '.docx': DocxExtractor()
        }
    
    def extract_text(self, file_path: str) -> tuple[str, Dict[str, Any]]:
        """
        Extract text from file using appropriate strategy
        
        Args:
            file_path: Path to the file to extract text from
            
        Returns:
            Tuple of (extracted_text, metadata)
            
        Raises:
            UnsupportedFileTypeError: If file type is not supported
            TextExtractionError: If extraction fails
            EmptyTextError: If no text is extracted
        """
        logger.info(f"Starting text extraction for file: {file_path}")
        
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            raise TextExtractionError(f"File not found: {file_path}")
        
        # Get file extension
        _, file_extension = os.path.splitext(file_path)
        file_extension = file_extension.lower()
        logger.debug(f"Detected file extension: {file_extension}")
        
        # Get appropriate extractor
        extractor = self.extractors.get(file_extension)
        if not extractor:
            logger.error(f"Unsupported file type: {file_extension}")
            raise UnsupportedFileTypeError(f"Unsupported file type: {file_extension}")
        
        logger.info(f"Using {extractor.__class__.__name__} for extraction")
        
        # Extract text and metadata
        try:
            text = extractor.extract(file_path)
            metadata = extractor.get_metadata(file_path)
            
            # Add common metadata
            metadata.update({
                "file_name": os.path.basename(file_path),
                "file_extension": file_extension,
                "character_count": len(text),
                "word_count": len(text.split()),
                "extraction_successful": True
            })
            
            logger.info(f"Text extraction completed successfully. Extracted {len(text)} characters, {len(text.split())} words")
            return text, metadata
            
        except Exception as e:
            logger.error(f"Text extraction failed for {file_path}: {str(e)}", exc_info=True)
            raise
    
    def get_supported_extensions(self) -> list[str]:
        """Get list of supported file extensions"""
        return list(self.extractors.keys())
    
    def is_supported(self, file_extension: str) -> bool:
        """Check if file extension is supported"""
        return file_extension.lower() in self.extractors 